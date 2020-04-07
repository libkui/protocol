import time

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 加粗
from docx.shared import Pt  # 磅数
from docx.oxml.ns import qn  # 中文格式
from docx.shared import Inches  # 图片尺寸
from docx.shared import RGBColor  # 颜色模块
from net_14_smtp.modules.syslog_bing import syslog_bing
import os


def create_word_for_syslog(db_name, add_img, save_word_name):
    today = time.strftime("%Y{y}%m{m}%d{d}", time.localtime()).format(y='年', m='月', d='日')

    document = Document()

    document.styles['Normal'].font.name = u'微软雅黑'
    document.styles['Normal'].font.size = Pt(14)
    # 设置文档的基础字体
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    # 设置文档的基础样式

    document.add_picture(add_img, width=Inches(6))
    # 在文件最上头插入图，宽度为6英寸

    p1 = document.add_paragraph()
    # 初始化建立第一个自然段
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 对齐方式为居中，没有这句的话默认左对齐。
    run1 = p1.add_run('乾颐堂Python强化班Syslog分析')
    run1.font.name = '微软雅黑'
    # 设置西文字体
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    # 设置中文字体
    run1.font.size = Pt(20)
    # 设置字体大小为20磅

    run1.font.bold = True
    # 设置加粗
    p1.space_after = Pt(1)
    # 段后距离5磅
    p1.space_before = Pt(5)
    # 段前距离5磅

    p2 = document.add_paragraph()
    syslog_result = syslog_bing(db_name, 'temp.png')

    run2 = p2.add_run('下面是最近一个小时的Syslog的数据统计! 显示排前三的Syslog严重级别与数量')
    run2.font.name = '仿宋_GB2312'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run2.font.size = Pt(10)
    table = document.add_table(rows=4, cols=3, style='Light Grid Accent 2')

    table.cell(0, 0).text = '严重级别'
    table.cell(0, 1).text = '数量'
    table.cell(0, 2).text = '百分比'

    total = sum([y for x, y in syslog_result])

    i = 1
    for x, y in syslog_result[:3]:
        table.cell(i, 0).text = x
        table.cell(i, 1).text = str(int(y))
        table.cell(i, 2).text = f'{(y/total)*100:.1f}'
        i += 1
    p3 = document.add_paragraph()
    run3 = p3.add_run('\r\n下面是最近一个小时的Syslog的数据统计饼状图分析!')
    run3.font.name = '仿宋_GB2312'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run3.font.size = Pt(10)
    document.add_picture('temp.png', width=Inches(3.0), height=Inches(3.0))
    os.remove('temp.png')
    document.save(save_word_name)


if __name__ == '__main__':
    create_word_for_syslog("../../net_9_syslog/practice_homework/syslog.sqlite", './src_img/logo.png', './saved_word/syslog-docx.docx')


